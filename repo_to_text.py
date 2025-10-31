#!/usr/bin/env python3
"""
GPL-3.0-or-later

repo_to_text.py — Flatten a directory/repository into a single text file for LLM ingestion.

Features
- Recursively walks a directory and concatenates text-like files into one output file.
- Skips obvious binaries (images, audio, video, archives) by extension and by content heuristic.
- Optional basic extraction for PDF (via PyMuPDF) and DOCX (via python-docx) if libraries are available.
- Adds clear section headers with relative paths and code fences with language tags where possible.
- Sensible default ignore list (e.g., .git, node_modules, __pycache__, venv, build artifacts).
- Streaming I/O (does not load everything into RAM). Handles large repos.

Usage
    python3 repo_to_text.py /path/to/repo -o repo_dump.txt

Tip
- Consider chunking the resulting file later if your LLM has context length limits.
"""
from __future__ import annotations

import argparse
import io
import os
import sys
import textwrap
from typing import Iterable, Optional, Tuple

# Optional imports for richer extraction
try:  # PDF text extraction
    import fitz  # PyMuPDF
    _HAS_PYMUPDF = True
except Exception:
    _HAS_PYMUPDF = False

try:  # DOCX text extraction
    import docx  # python-docx
    _HAS_PYTHON_DOCX = True
except Exception:
    _HAS_PYTHON_DOCX = False

# --- Configuration helpers ---

DEFAULT_IGNORED_DIRS = {
    ".git", ".hg", ".svn", ".bzr",
    "node_modules", "__pycache__", ".pytest_cache", ".mypy_cache",
    "venv", ".venv", "env", ".env", "ENV",
    "dist", "build", "out", "target", "coverage", "site-packages",
    ".idea", ".vscode", ".DS_Store"
}

# Obvious binary or non-text by extension (skip always)
ALWAYS_SKIP_EXTS = {
    # images
    ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".webp", ".svgz",
    # audio/video
    ".mp3", ".wav", ".flac", ".ogg", ".mp4", ".mkv", ".avi", ".mov", ".webm",
    # archives/binaries
    ".zip", ".tar", ".gz", ".tgz", ".bz2", ".xz", ".7z", ".rar",
    ".exe", ".dll", ".so", ".dylib", ".class", ".o", ".a", ".bin", ".wasm",
    ".pdf"  # handled separately if PyMuPDF is available
}

# Extensions we strongly expect to be text (use direct read)
PREFER_TEXT_EXTS = {
    # code
    ".py", ".pyi", ".pyx", ".pxd",
    ".js", ".jsx", ".ts", ".tsx",
    ".java", ".kt", ".kts", ".scala",
    ".rb", ".go", ".rs", ".c", ".h", ".cpp", ".hpp", ".cc", ".hh", ".m", ".mm",
    ".cs", ".vb", ".fs", ".fsx",
    ".php", ".r", ".jl", ".pl", ".lua", ".sh", ".bash", ".zsh", ".fish",
    ".ps1", ".psm1", ".psd1", ".sql", ".yaml", ".yml", ".json", ".jsonc",
    ".toml", ".ini", ".cfg", ".conf", ".env", ".properties",
    ".html", ".htm", ".xml", ".svg", ".xhtml", ".xsd",
    ".md", ".rst", ".txt", ".csv", ".tsv", ".log", ".tex",
    ".gradle", ".sbt", ".cabal", ".nim", ".hs"
}

# Allow optional lightweight conversion for docx if lib exists
DOCX_EXTS = {".docx"}
PDF_EXTS = {".pdf"}

# Map extensions to language tags for code fences
EXT_TO_LANG = {
    ".py": "python", ".js": "javascript", ".jsx": "jsx", ".ts": "typescript",
    ".tsx": "tsx", ".java": "java", ".kt": "kotlin", ".kts": "kotlin",
    ".scala": "scala", ".rb": "ruby", ".go": "go", ".rs": "rust",
    ".c": "c", ".h": "c", ".cpp": "cpp", ".hpp": "cpp", ".cc": "cpp", ".hh": "cpp",
    ".m": "objective-c", ".mm": "objective-cpp",
    ".cs": "csharp", ".vb": "vbnet", ".fs": "fsharp", ".fsx": "fsharp",
    ".php": "php", ".r": "r", ".jl": "julia", ".pl": "perl", ".lua": "lua",
    ".sh": "bash", ".bash": "bash", ".zsh": "zsh", ".fish": "fish",
    ".ps1": "powershell", ".psm1": "powershell", ".psd1": "powershell",
    ".sql": "sql", ".yaml": "yaml", ".yml": "yaml", ".json": "json",
    ".toml": "toml", ".ini": "ini", ".cfg": "ini", ".conf": "ini",
    ".html": "html", ".htm": "html", ".xml": "xml", ".svg": "xml",
    ".md": "markdown", ".rst": "rst", ".txt": "text", ".csv": "csv", ".tsv": "tsv",
    ".tex": "tex"
}

# --- Utility functions ---

def is_probably_binary(sample: bytes) -> bool:
    """Heuristic: treat as binary if NUL present or too many non-text bytes.

    We allow common whitespace and printable ASCII, plus a bit of extended bytes.
    """
    if not sample:
        return False
    # Any NUL byte is a strong binary signal
    if b"\x00" in sample:
        return True
    # Count "suspicious" control chars (exclude common whitespace \t,\n,\r,\f)
    text_whitelist = set(b"\t\n\r\f\b\x0c")  # include form feed for safety
    suspicious = 0
    for b in sample:
        if b == 0:
            return True
        if 32 <= b <= 126:  # printable ASCII
            continue
        if b in text_whitelist:
            continue
        if 127 <= b <= 159:  # C1 controls often binary-ish
            suspicious += 1
            continue
        # Extended bytes: allow some, but track density
        # Count as suspicious; many of them likely not plain text in mixed repos
        suspicious += 1
    # If more than 30% of bytes are suspicious, call it binary
    return suspicious / max(1, len(sample)) > 0.30


def iter_files(root: str, include_hidden: bool, ignored_dirs: set[str]) -> Iterable[str]:
    for dirpath, dirnames, filenames in os.walk(root, followlinks=False):
        # Filter directories in-place
        keep_dirs = []
        for d in dirnames:
            if not include_hidden and d.startswith('.'):
                continue
            if d in ignored_dirs:
                continue
            keep_dirs.append(d)
        dirnames[:] = keep_dirs

        for fname in filenames:
            if not include_hidden and fname.startswith('.'):
                continue
            yield os.path.join(dirpath, fname)


def safe_relpath(path: str, start: str) -> str:
    try:
        return os.path.relpath(path, start)
    except Exception:
        return path


def read_text_file(path: str, encodings: Tuple[str, ...] = ("utf-8", "utf-8-sig", "latin-1", "cp1252")) -> Optional[str]:
    for enc in encodings:
        try:
            with io.open(path, 'r', encoding=enc, errors='strict') as f:
                return f.read()
        except Exception:
            continue
    # last resort: replace errors with placeholders to avoid crash
    try:
        with io.open(path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    except Exception:
        return None


def extract_docx_text(path: str) -> Optional[str]:
    if not _HAS_PYTHON_DOCX:
        return None
    try:
        document = docx.Document(path)
        parts = []
        for para in document.paragraphs:
            parts.append(para.text)
        # tables
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                parts.append("\t".join(cells))
        return "\n".join(parts)
    except Exception:
        return None


def extract_pdf_text(path: str) -> Optional[str]:
    if not _HAS_PYMUPDF:
        return None
    try:
        doc = fitz.open(path)
        parts = []
        for page in doc:
            parts.append(page.get_text())
        doc.close()
        return "\n".join(parts)
    except Exception:
        return None


def fence_for_ext(ext: str) -> Tuple[str, str]:
    lang = EXT_TO_LANG.get(ext.lower(), "")
    if lang:
        return f"```{lang}\n", "\n```\n"
    return "```\n", "\n```\n"


# --- Core processing ---

def process_file(path: str, root: str, args) -> Optional[str]:
    """Return a string chunk to append to output, or None to skip.

    Uses streaming-ish approach by returning one big string per file to keep code simple.
    The caller writes to output file incrementally.
    """
    ext = os.path.splitext(path)[1].lower()

    # Skip by extension if clearly non-text, except allow PDFs/DOCX if we can extract
    if ext in ALWAYS_SKIP_EXTS:
        # allow PDF via special case if extractor exists
        if ext in PDF_EXTS and args.pdf and _HAS_PYMUPDF:
            pass
        else:
            return None

    if (args.only_ext and ext not in args.only_ext) or (args.skip_ext and ext in args.skip_ext):
        return None

    try:
        size = os.path.getsize(path)
    except Exception:
        size = -1

    if args.max_file_size and size > args.max_file_size:
        return None

    rel = safe_relpath(path, root)

    # Decide handling path
    content: Optional[str] = None

    if ext in DOCX_EXTS and args.docx:
        content = extract_docx_text(path)

    if content is None and ext in PDF_EXTS and args.pdf:
        content = extract_pdf_text(path)

    if content is None:
        # For plain text-like files, quick binary sniff first
        try:
            with open(path, 'rb') as fb:
                head = fb.read(4096)
        except Exception:
            return None
        # If extension not in PREFER_TEXT_EXTS, use binary heuristic
        if ext not in PREFER_TEXT_EXTS and is_probably_binary(head):
            return None
        # Try reading as text
        content = read_text_file(path)

    if content is None:
        return None

    # Normalize newlines
    content = content.replace('\r\n', '\n').replace('\r', '\n')

    # Compose header and fenced block
    header = f"\n===== FILE: {rel} (size={size} bytes) =====\n"
    open_fence, close_fence = fence_for_ext(ext)

    return header + open_fence + content + close_fence


def write_header(out, root: str, args):
    meta = textwrap.dedent(f"""
    ==============================================
    Repo to Text Export
    Root: {os.path.abspath(root)}
    Generated: (local time) {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    Include hidden: {args.include_hidden}
    Max file size (bytes): {args.max_file_size if args.max_file_size else 'None'}
    PDF extraction enabled: {args.pdf and _HAS_PYMUPDF}
    DOCX extraction enabled: {args.docx and _HAS_PYTHON_DOCX}
    Ignored dirs: {', '.join(sorted(args.ignore_dirs))}
    Only extensions: {', '.join(sorted(args.only_ext)) if args.only_ext else 'None'}
    Skipped extensions: {', '.join(sorted(args.skip_ext)) if args.skip_ext else 'None'}
    ==============================================
    """)
    out.write(meta + "\n")


def main(argv: Optional[Iterable[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="Flatten a repository into a single text file.")
    parser.add_argument("root", help="Root directory to scan")
    parser.add_argument("-o", "--output", default="repo_dump.txt", help="Output text file path")
    parser.add_argument("--include-hidden", action="store_true", help="Include hidden files and directories")
    parser.add_argument("--no-default-ignores", action="store_true", help="Do not use the built-in ignored directories")
    parser.add_argument("--ignore-dir", action="append", default=[], help="Additional directory names to ignore (can repeat)")
    parser.add_argument("--max-file-size", type=int, default=0, help="Skip files larger than this many bytes (0 = no limit)")
    parser.add_argument("--pdf", action="store_true", help="Attempt to extract text from PDFs (requires PyMuPDF)")
    parser.add_argument("--docx", action="store_true", help="Attempt to extract text from DOCX (requires python-docx)")
    parser.add_argument("--only-ext", nargs="*", default=[], help="Only include files with these extensions (e.g. .py .md)")
    parser.add_argument("--skip-ext", nargs="*", default=[], help="Skip files with these extensions (e.g. .log .csv)")

    args = parser.parse_args(argv)

    root = os.path.abspath(args.root)
    if not os.path.isdir(root):
        print(f"Error: '{root}' is not a directory", file=sys.stderr)
        return 2

    ignore_dirs = set(args.ignore_dir)
    if not args.no_default_ignores:
        ignore_dirs |= set(DEFAULT_IGNORED_DIRS)
    args.ignore_dirs = ignore_dirs  # attach to args for header printing

    # Normalize extensions to dot-lower format
    args.only_ext = set(e if e.startswith('.') else f'.{e}' for e in args.only_ext)
    args.only_ext = set(e.lower() for e in args.only_ext)
    args.skip_ext = set(e if e.startswith('.') else f'.{e}' for e in args.skip_ext)
    args.skip_ext = set(e.lower() for e in args.skip_ext)

    if args.max_file_size is not None and args.max_file_size <= 0:
        args.max_file_size = 0

    count_total = 0
    count_included = 0

    with io.open(args.output, 'w', encoding='utf-8') as out:
        write_header(out, root, args)
        for path in iter_files(root, include_hidden=args.include_hidden, ignored_dirs=ignore_dirs):
            count_total += 1
            chunk = process_file(path, root, args)
            if chunk is not None:
                out.write(chunk)
                if not chunk.endswith('\n'):
                    out.write('\n')
                count_included += 1

        # Trailer summary
        out.write(f"\n===== SUMMARY =====\n")
        out.write(f"Files scanned: {count_total}\n")
        out.write(f"Files included: {count_included}\n")

    print(f"Wrote {args.output}. Files scanned: {count_total}, included: {count_included}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
